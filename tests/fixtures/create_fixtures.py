"""
Generate test fixture .docx files.
Run once before running the test suite:
    python tests/fixtures/create_fixtures.py
"""
import zipfile
import tempfile
from pathlib import Path
from docx import Document
from lxml import etree

FIXTURES_DIR = Path(__file__).parent
WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = f"{{{WNS}}}"


def _repack(base_path: Path, doc_xml_bytes: bytes, output_path: Path,
            comments_xml_bytes: bytes = None):
    """Unzip base, swap document.xml (and optionally comments.xml), repack."""
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        with zipfile.ZipFile(base_path) as z:
            z.extractall(tmp)

        (tmp / "word" / "document.xml").write_bytes(doc_xml_bytes)

        if comments_xml_bytes is not None:
            (tmp / "word" / "comments.xml").write_bytes(comments_xml_bytes)

        output_path.unlink(missing_ok=True)
        with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for f in sorted(tmp.rglob("*")):
                if f.is_file():
                    zout.write(f, f.relative_to(tmp))


def create_clean_sow():
    """Standard SOW with no tracked changes or comments."""
    doc = Document()
    doc.add_heading("Statement of Work", level=1)
    doc.add_paragraph(
        "This is a standard statement of work with no tracked changes or comments."
    )
    doc.add_heading("Scope of Work", level=2)
    doc.add_paragraph("The consultant will deliver the following services.")
    doc.add_paragraph("Service delivery will occur over a 12-week period.")
    path = FIXTURES_DIR / "clean_sow.docx"
    doc.save(str(path))
    print(f"Created: {path}")


def create_tracked_change_proposal():
    """
    Proposal with tracked insertions and deletions.
    Clean text: "This is newly inserted sample text."
    Deleted text: "old text"
    """
    doc = Document()
    doc.add_heading("Project Proposal", level=1)
    doc.add_paragraph("PLACEHOLDER")
    base_path = FIXTURES_DIR / "_base_tc.docx"
    doc.save(str(base_path))

    with zipfile.ZipFile(base_path) as z:
        doc_xml = z.read("word/document.xml")

    root = etree.fromstring(doc_xml)
    body = root.find(f".//{W}body")

    # Remove the placeholder paragraph
    for p in list(body.findall(f"{W}p")):
        texts = "".join(t.text or "" for t in p.findall(f".//{W}t"))
        if "PLACEHOLDER" in texts:
            body.remove(p)
            break

    # Inject a paragraph with tracked insertion and deletion
    # Final accepted text: "This is newly inserted sample text."
    tc_para_xml = (
        f'<w:p xmlns:w="{WNS}">'
        f'<w:r><w:t xml:space="preserve">This is </w:t></w:r>'
        f'<w:ins w:id="1" w:author="Jane Reviewer" w:date="2026-01-01T00:00:00Z">'
        f'<w:r><w:t xml:space="preserve">newly inserted </w:t></w:r>'
        f'</w:ins>'
        f'<w:del w:id="2" w:author="Jane Reviewer" w:date="2026-01-01T00:00:00Z">'
        f'<w:r><w:delText>old text </w:delText></w:r>'
        f'</w:del>'
        f'<w:r><w:t>sample text.</w:t></w:r>'
        f'</w:p>'
    )
    tc_para = etree.fromstring(tc_para_xml)

    sect = body.find(f"{W}sectPr")
    if sect is not None:
        body.insert(list(body).index(sect), tc_para)
    else:
        body.append(tc_para)

    doc_xml_bytes = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    output_path = FIXTURES_DIR / "tracked_change_proposal.docx"
    _repack(base_path, doc_xml_bytes, output_path)
    base_path.unlink()
    print(f"Created: {output_path}")


def create_table_comment_doc():
    """
    Document with a table and reviewer comments.
    Includes: one actionable comment, one question comment.
    """
    doc = Document()
    doc.add_heading("Deliverables", level=1)
    doc.add_paragraph("The consultant shall complete all assigned tasks.")
    doc.add_paragraph("Please reword this sentence.")  # comment anchor

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Role"
    table.cell(0, 1).text = "Responsibility"
    table.cell(1, 0).text = "Consultant"
    table.cell(1, 1).text = "Deliver the work"

    base_path = FIXTURES_DIR / "_base_comment.docx"
    doc.save(str(base_path))

    with zipfile.ZipFile(base_path) as z:
        doc_xml = z.read("word/document.xml")

    root = etree.fromstring(doc_xml)
    body = root.find(f".//{W}body")

    # Find and replace the "Please reword" paragraph with one that has comment markers
    for p in list(body.findall(f".//{W}p")):
        texts = "".join(t.text or "" for t in p.findall(f".//{W}t"))
        if "Please reword this sentence." in texts:
            parent = p.getparent()
            idx = list(parent).index(p)
            parent.remove(p)

            new_para_xml = (
                f'<w:p xmlns:w="{WNS}">'
                f'<w:commentRangeStart w:id="1"/>'
                f'<w:r><w:t>Please reword this sentence.</w:t></w:r>'
                f'<w:commentRangeEnd w:id="1"/>'
                f'<w:r>'
                f'<w:rPr><w:rStyle w:val="CommentReference"/></w:rPr>'
                f'<w:commentReference w:id="1"/>'
                f'</w:r>'
                f'</w:p>'
            )
            parent.insert(idx, etree.fromstring(new_para_xml))
            break

    doc_xml_bytes = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    comments_xml = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<w:comments xmlns:w="{WNS}">'
        f'<w:comment w:id="1" w:author="Alice Reviewer" w:date="2026-01-15T10:00:00Z">'
        f'<w:p><w:r><w:t>Please reword this sentence.</w:t></w:r></w:p>'
        f'</w:comment>'
        f'</w:comments>'
    ).encode("utf-8")

    output_path = FIXTURES_DIR / "table_comment_doc.docx"
    _repack(base_path, doc_xml_bytes, output_path, comments_xml)
    base_path.unlink()
    print(f"Created: {output_path}")


if __name__ == "__main__":
    create_clean_sow()
    create_tracked_change_proposal()
    create_table_comment_doc()
    print("\nAll fixtures created.")
